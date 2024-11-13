import os
import streamlit as st
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import time
from dotenv import load_dotenv
import tempfile
from docx import Document as DocxDocument
import textract
from langchain.schema.document import Document

# Load environment variables for API keys
load_dotenv()

# Load GROQ and Google API keys
groq_api_key = os.getenv("GROQ_API_KEY")
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# Set up authentication credentials (you can replace these with environment variables for security)
AUTH_USERNAME = "admin"
AUTH_PASSWORD = "password"

# Initialize Streamlit session state for authentication
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

# Authentication function
def authenticate(username, password):
    return username == AUTH_USERNAME and password == AUTH_PASSWORD

# Login page
def login():
    st.title("Login to Access RAG Chatbot")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    
    if st.button("Login"):
        if authenticate(username, password):
            st.session_state.authenticated = True
            st.success("Login successful! Access granted.")
        else:
            st.error("Invalid username or password. Please try again.")

# Run login page if not authenticated
if not st.session_state.authenticated:
    login()
else:
    # Streamlit App Title after login
    st.title("RAG CHATBOT with Caching and Authentication")

    # Initialize the Groq Language Model
    llm = ChatGroq(groq_api_key=groq_api_key, model_name="Gemma-7b-it")

    # Set up the prompt template
    prompt_template = ChatPromptTemplate.from_template(
        """
        Answer the questions based on the provided context only.
        Provide the most accurate response based on the question.

        <context>
        {context}
        <context>
        Question: {input}
        """
    )

    # Function to read content from DOCX files
    def read_docx(file):
        doc = DocxDocument(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to read content from DOC files
    def read_doc(file):
        return textract.process(file, encoding='utf-8').decode('utf-8')

    # Initialize vector embeddings
    def vector_embedding(uploaded_files):
        if "vectors" not in st.session_state:
            st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
            all_docs = []  # Store all documents from all files

            for uploaded_file in uploaded_files:
                file_extension = uploaded_file.name.split(".")[-1].lower()
                
                if file_extension == "pdf":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    pdf_loader = PyPDFLoader(temp_path)
                    pdf_docs = pdf_loader.load()
                    all_docs.extend([Document(page_content=doc.page_content, metadata=doc.metadata) for doc in pdf_docs])
                    os.remove(temp_path)

                elif file_extension == "docx":
                    text = read_docx(uploaded_file)
                    all_docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))

                elif file_extension == "doc":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    text = read_doc(temp_path)
                    all_docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
                    os.remove(temp_path)

            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            st.session_state.final_documents = text_splitter.split_documents(all_docs)
            
            # Create vector store using FAISS
            st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)
            st.write("Vector Store DB is Ready!")

    # Drag-and-drop PDF, DOC, and DOCX upload
    uploaded_files = st.file_uploader("Drag and drop PDF, DOC, or DOCX files here", type=["pdf", "doc", "docx"], accept_multiple_files=True)

    # Button to create vector store if files are uploaded
    if uploaded_files and st.button("Create Vector Store"):
        vector_embedding(uploaded_files)

    # Initialize cache for previous queries
    if "cache" not in st.session_state:
        st.session_state.cache = {}

    # Interface for question input
    user_query = st.text_input("Enter your question based on the documents:")

    # Process and display the response based on the user's query
    if user_query:
        # Check cache for the answer
        if user_query in st.session_state.cache:
            st.write("Cached Response:", st.session_state.cache[user_query]["answer"])
            with st.expander("Document Similarity Search Results from Cache"):
                for i, doc in enumerate(st.session_state.cache[user_query]["context"]):
                    st.write(f"Relevant Chunk {i+1} (Source: {doc.metadata['source']}):")
                    st.write(doc.page_content)
                    st.write("-----------------------")
        else:
            if "vectors" in st.session_state:
                document_chain = create_stuff_documents_chain(llm, prompt_template)
                retriever = st.session_state.vectors.as_retriever()
                retrieval_chain = create_retrieval_chain(retriever, document_chain)
                
                start = time.process_time()
                
                response = retrieval_chain.invoke({"input": user_query})
                st.write("Response:", response["answer"])
                
                with st.expander("Document Similarity Search Results"):
                    for i, doc in enumerate(response["context"]):
                        st.write(f"Relevant Chunk {i+1} (Source: {doc.metadata['source']}):")
                        st.write(doc.page_content)
                        st.write("-----------------------")
                
                # Cache the answer and context
                st.session_state.cache[user_query] = {
                    "answer": response["answer"],
                    "context": response["context"]
                }
                
                end = time.process_time()
                st.write(f"Time taken for response: {end - start} seconds")
            else:
                st.warning("Please create the vector store first by clicking the button above.")
