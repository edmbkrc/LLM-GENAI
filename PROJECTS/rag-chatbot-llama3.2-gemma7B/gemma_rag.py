import os
import streamlit as st
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import ChatMessage
import tempfile
from docx import Document as DocxDocument
import textract
from langchain.schema.document import Document
from dotenv import load_dotenv

def gemma_rag_main():
    # Load environment variables for API keys
    load_dotenv()

    # Load GROQ and Google API keys
    groq_api_key = os.getenv("GROQ_API_KEY")
    os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

    # Initialize ChatGroq model
    llm = ChatGroq(groq_api_key=groq_api_key, 
                   model_name="Gemma-7b-it",
                   max_tokens=900
                   )

    # Prompt Template
    prompt_template = """
    Given the following context, answer the question as accurately and comprehensively as possible.
    Combine relevant information from all provided chunks if necessary. 
    If the answer cannot be derived from the context, respond only with: "Sorry, I can only answer questions based on the provided documents."
    Context:
    {context}

    Question:
    {input}
    """

    # Helper Functions for Reading Files
    def read_docx(file):
        doc = DocxDocument(file)
        return "\n".join([para.text for para in doc.paragraphs])

    def read_doc(file):
        return textract.process(file, encoding="utf-8").decode("utf-8")

    # Update Vector Store
    def update_vector_store(uploaded_files):
        new_docs = []

        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name  # Extract file name
            if file_name in st.session_state.processed_files:
                st.write(f"File '{file_name}' already processed. Skipping.")
                continue

            file_extension = file_name.split(".")[-1].lower()

            try:
                if file_extension == "pdf":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    pdf_loader = PyPDFLoader(temp_path)
                    pdf_docs = pdf_loader.load()
                    new_docs.extend([
                        Document(page_content=doc.page_content, 
                                 metadata={"source": file_name, "page": doc.metadata.get("page", "Unknown")})
                        for doc in pdf_docs
                    ])
                    os.remove(temp_path)

                elif file_extension == "docx":
                    text = read_docx(uploaded_file)
                    if text.strip():
                        new_docs.append(Document(page_content=text, metadata={"source": file_name}))

                elif file_extension == "doc":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    text = read_doc(temp_path)
                    if text.strip():
                        new_docs.append(Document(page_content=text, metadata={"source": file_name}))
                    os.remove(temp_path)

                st.session_state.processed_files.add(file_name)
            except Exception as e:
                st.error(f"Error processing file {file_name}: {e}")

        # Split Documents into Chunks and Update Vectorstore
        if new_docs:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
            new_chunks = text_splitter.split_documents(new_docs)

            if not new_chunks:
                st.warning("No valid content found in the uploaded files. Skipping.")
                return

            texts = [chunk.page_content for chunk in new_chunks]
            metadatas = [chunk.metadata for chunk in new_chunks]
            embeddings = st.session_state.embeddings.embed_documents(texts)

            if st.session_state.vectors is not None:
                st.session_state.vectors.add_texts(texts=texts, embeddings=embeddings, metadatas=metadatas)
                st.write("Vectorstore updated with new chunks!")
            else:
                st.session_state.vectors = FAISS.from_texts(texts=texts, embedding=st.session_state.embeddings, metadatas=metadatas)
                st.write("Vectorstore created successfully!")

    # Initialize Session State
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = set()
    if "embeddings" not in st.session_state:
        st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    if "vectors" not in st.session_state:
        st.session_state.vectors = None
    if "cache" not in st.session_state:
        st.session_state.cache = {}
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Streamlit UI
    st.title("RAG Chatbot with Gemma")

    # File Upload
    uploaded_files = st.file_uploader(
        "Upload your documents (PDF, DOC, DOCX)", type=["pdf", "doc", "docx"], accept_multiple_files=True
    )

    if uploaded_files:
        with st.spinner("Processing files..."):
            update_vector_store(uploaded_files)

    # Chat Interface
    st.subheader("Chat with your documents")
    user_input = st.text_input("You:", key="input")

    if user_input:
        # Check for Cache
        if user_input in st.session_state.cache:
            cached_response = st.session_state.cache[user_input]
            st.write("**Cached Response:**")
            st.write(f"**You:** {user_input}")
            st.write(f"**Bot:** {cached_response['answer']}")
            with st.expander("Retrieved Chunks"):
                for i, doc in enumerate(cached_response["retrieved_docs"], start=1):
                    source = doc.metadata.get("source", "Unknown")
                    page = doc.metadata.get("page", "Unknown")
                    st.write(f"**Chunk {i} (Source: {source}, Page: {page}):**")
                    st.write(doc.page_content)
                    st.write("---")
        else:
            if st.session_state.vectors is None:
                st.warning("Vectorstore is not initialized. Please upload documents and create the vectorstore first.")
            else:
                retriever = st.session_state.vectors.as_retriever(search_kwargs={"k": 5, "distance_threshold": 0.7})
                retrieved_docs = retriever.get_relevant_documents(user_input)

                if retrieved_docs:
                    context = "\n".join([doc.page_content for doc in retrieved_docs])
                    formatted_prompt = prompt_template.format(context=context, input=user_input)

                    # Format Query for LLM
                    messages = [ChatMessage(role="user", content=formatted_prompt)]
                    response = llm(messages).content

                    # Cache and Display Response
                    st.session_state.cache[user_input] = {"answer": response, "retrieved_docs": retrieved_docs}
                    st.session_state.messages.append((user_input, response))

                    st.write(f"**You:** {user_input}")
                    st.write(f"**Bot:** {response}")

                    with st.expander("Retrieved Chunks"):
                        for i, doc in enumerate(retrieved_docs, start=1):
                            source = doc.metadata.get("source", "Unknown")
                            page = doc.metadata.get("page", "Unknown")
                            st.write(f"**Chunk {i} (Source: {source},Page {str(page) if isinstance(page, int) else page}):**")
                            st.write(doc.page_content)
                            st.write("---")
                else:
                    st.write("Sorry, I could not find relevant information in the uploaded documents.")

    # Chat History
    with st.expander("Chat History"):
        for user_msg, bot_msg in st.session_state.messages:
            st.write(f"**You:** {user_msg}")
            st.write(f"**Bot:** {bot_msg}")
