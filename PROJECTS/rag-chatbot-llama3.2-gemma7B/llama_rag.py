import os
import tempfile
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
import textract
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama
from langchain_core.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.schema.document import Document


def llama_rag_main():

    # Function to read DOCX files
    def read_docx(file):
        doc = DocxDocument(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to read DOC files
    def read_doc(file):
        return textract.process(file, encoding="utf-8").decode("utf-8")

    # Function to update vector store dynamically
    def update_vector_store(uploaded_files):
        new_docs = []
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            if file_name in st.session_state.processed_files:
                st.write(f"File '{file_name}' already processed. Skipping.")
                continue

            file_extension = file_name.split(".")[-1].lower()
            try:
                if file_extension == "pdf":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    pdf_loader = PyPDFLoader(temp_path)
                    pdf_docs = pdf_loader.load()
                    new_docs.extend([
                        Document(page_content=doc.page_content, metadata={"source": file_name})
                        for doc in pdf_docs
                    ])
                    os.remove(temp_path)

                elif file_extension == "docx":
                    text = read_docx(uploaded_file)
                    if text.strip():
                        new_docs.append(Document(page_content=text, metadata={"source": file_name}))

                elif file_extension == "doc":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_path = temp_file.name
                    text = read_doc(temp_path)
                    if text.strip():
                        new_docs.append(Document(page_content=text, metadata={"source": file_name}))
                    os.remove(temp_path)

                st.session_state.processed_files.add(file_name)

            except Exception as e:
                st.error(f"Error processing file {file_name}: {e}")

        # Split documents into chunks and update vectorstore
        if new_docs:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
            new_chunks = text_splitter.split_documents(new_docs)

            if not new_chunks:
                st.warning("No valid content found in the uploaded files. Skipping.")
                return
            # Extracting Text and Metadata and embeddings
            texts = [chunk.page_content for chunk in new_chunks]
            metadatas = [chunk.metadata for chunk in new_chunks]
            embeddings = st.session_state.embeddings.embed_documents(texts)
            # updates an existing vector store or creates a new one depending on the current state.
            if st.session_state.vectors:
                st.session_state.vectors.add_texts(texts=texts, embeddings=embeddings, metadatas=metadatas)
                st.write("Vectorstore updated successfully!")
            else:
                st.session_state.vectors = FAISS.from_texts(texts=texts, embedding=st.session_state.embeddings, metadatas=metadatas)
                st.write("Vectorstore created successfully!")

    # initializes various components in the Streamlit session state 
    # to ensure that essential data structures and models are properly 
    # This ensures that critical objects, like models or cached data, 
    # remain available without being reloaded or recalculated repeatedly.
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = set() # for storing unique file names, allowing fast lookups, additions, and deletions.
    if "embeddings" not in st.session_state:
        st.session_state.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L12-v2")
    if "vectors" not in st.session_state:
        st.session_state.vectors = None # The vector store is created or updated dynamically during runtime when data is uploaded and processed.
    if "cache" not in st.session_state:
        st.session_state.cache = {} # dictionary for caching intermediate data or results, such as partially processed files or embeddings.
    if "messages" not in st.session_state:
        st.session_state.messages = [] 

    # Streamlit UI
    st.title("RAG Chatbot with Llama")
    # User-configurable parameters
    st.sidebar.header("Model Parameters")
    prompt_prefix = st.sidebar.text_area("Prompt Prefix", """
    Given the following context, answer the question as accurately and comprehensively as possible.
    Combine relevant information from all provided chunks if necessary. 
    If the answer cannot be derived from the context, respond only with: "Sorry, I can only answer questions based on the provided documents."
    """)
    
    temperature = st.sidebar.slider("Temperature", 0.0, 1.0, 0.7)
    top_p = st.sidebar.slider("Top-p", 0.0, 1.0, 0.9)
    top_k = st.sidebar.slider("Top-k", 1, 100, 40)
    

    # File uploader
    uploaded_files = st.file_uploader(
        "Upload your documents (PDF, DOC, DOCX)", type=["pdf", "doc", "docx"], accept_multiple_files=True
    )

    if uploaded_files:
        with st.spinner("Processing files..."):
            update_vector_store(uploaded_files)

    # Chat interface
    st.subheader("Chat with your documents")
    user_input = st.text_input("You:", key="input")

    if user_input:
        # Ensure vectorstore is initialized
        if not st.session_state.vectors:
            st.warning("Vectorstore is not initialized. Please upload documents and create the vectorstore first.")
            return

        # Check if the query exists in the cache
        if user_input in st.session_state.cache:
            cached_response = st.session_state.cache[user_input]
            st.write("**Cached Response:** This response is retrieved from the cache.")
            st.write(f"**You:** {user_input}")
            st.write(f"**Bot:** {cached_response['answer']}")
            with st.expander("Retrieved Chunks"):
                for i, doc in enumerate(cached_response["retrieved_docs"], start=1):
                    st.write(f"**Chunk {i} (Source: {doc.metadata.get('source', 'Unknown')}):**")
                    st.write(doc.page_content)
                    st.write("---")
        else:
            retriever = st.session_state.vectors.as_retriever(search_kwargs={"k": 5, "distance_threshold": 0.7})
            retrieved_docs = retriever.get_relevant_documents(user_input)

            if retrieved_docs:
                context = "\n".join([doc.page_content for doc in retrieved_docs])
                formatted_prompt = f"""
                {prompt_prefix}

                Context:
                {context}

                Question:
                {user_input}

                Answer:
                """

                llm = Ollama(
                    model="llama3.2:1b",
                    callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
                    temperature = temperature,
                    top_p = top_p,
                    top_k = top_k
                    )
                response = llm(formatted_prompt)
            else:
                response = "Sorry, I can only answer questions based on the provided documents. No relevant information was found."

            st.session_state.cache[user_input] = {"answer": response, "retrieved_docs": retrieved_docs}
            st.session_state.messages.append((user_input, response))

            st.write(f"**You:** {user_input}")
            st.write(f"**Bot:** {response}")

            with st.expander("Retrieved Chunks"):
                for i, doc in enumerate(retrieved_docs, start=1):
                    source = doc.metadata.get("source", "Unknown")
                    page = doc.metadata.get("page", "Unknown")
                    st.write(f"**Chunk {i} (Source: {source},Page {str(page) if isinstance(page, int) else page}):**")
                    st.write(doc.page_content)
                    st.write("---")
    # Chat history
    with st.expander("Chat History"):
        for user_msg, bot_msg in st.session_state.messages:
            st.write(f"**You:** {user_msg}")
            st.write(f"**Bot:** {bot_msg}")
